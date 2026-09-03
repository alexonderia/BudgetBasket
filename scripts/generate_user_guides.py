import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
CONTENT_PATH = ROOT / "frontend" / "src" / "content" / "userGuideContent.json"
OUTPUT_PATH = ROOT / "docs" / "user-guides" / "BudgetBasket_User_Guide.docx"
ROLE_ORDER = ("employee", "economist", "approver", "zgd", "admin")
ROLE_COLORS = {
    "employee": "2C5F7C",
    "economist": "2C5F7C",
    "approver": "2C5F7C",
    "zgd": "2C5F7C",
    "admin": "2C5F7C",
}


def load_content() -> dict:
    with CONTENT_PATH.open("r", encoding="utf-8") as source:
        return json.load(source)


def keep_with_next(paragraph) -> None:
    properties = paragraph._p.get_or_add_pPr()
    properties.append(OxmlElement("w:keepNext"))


def shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(1.7)
    add_page_number(section.footer.paragraphs[0])
    header = section.header.paragraphs[0]
    header.text = "BUDGETBASKET  ·  РУКОВОДСТВО ПОЛЬЗОВАТЕЛЯ"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("6B7D8A")

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color in (
        ("Title", 25, "173B57"),
        ("Heading 1", 17, "173B57"),
        ("Heading 2", 13, "2C5F7C"),
        ("Heading 3", 11, "2C5F7C"),
    ):
        style = document.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def add_cover(document: Document, content: dict) -> None:
    document.add_paragraph()
    title = document.add_heading(content["title"], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph("Полное руководство по текущей версии приложения")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(73, 93, 109)

    document.add_paragraph()
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Версия {content['version']}\n").bold = True
    meta.add_run(f"Актуально на {content['updated']}")

    document.add_paragraph()
    intro = document.add_paragraph(content["intro"])
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    document.add_paragraph()
    usage = document.add_paragraph()
    usage.add_run("Как пользоваться. ").bold = True
    usage.add_run("Сначала ознакомьтесь с общей частью, затем перейдите к разделу своей роли. В каждом сценарии действия, результат и следующий шаг приведены в одном порядке.")
    usage.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    document.add_paragraph()
    add_journey(document, content["journey"])

    document.add_page_break()


def add_journey(document: Document, stages: list[dict]) -> None:
    heading = document.add_heading("Как движется бюджет", level=2)
    keep_with_next(heading)
    table = document.add_table(rows=2, cols=len(stages))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for index, stage in enumerate(stages):
        title_cell = table.cell(0, index)
        detail_cell = table.cell(1, index)
        shade_cell(title_cell, "2C5F7C")
        shade_cell(detail_cell, "EDF4F8")
        title_cell.text = f"{index + 1}. {stage['title']}"
        detail_cell.text = stage["detail"]
        for cell in (title_cell, detail_cell):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=90, start=85, bottom=90, end=85)
        title = title_cell.paragraphs[0]
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)
        detail = detail_cell.paragraphs[0]
        detail.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in detail.runs:
            run.font.size = Pt(8)


def add_quick_start(document: Document, role_label: str, steps: list[str], color: str) -> None:
    heading = document.add_heading(f"Быстрый маршрут: {role_label}", level=2)
    keep_with_next(heading)
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, step in enumerate(steps):
        cells = table.add_row().cells
        cells[0].width = Cm(1)
        cells[1].width = Cm(16)
        cells[0].text = str(index + 1)
        cells[1].text = step
        shade_cell(cells[0], color)
        shade_cell(cells[1], "F3F7FA")
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=75, start=110, bottom=75, end=110)
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        for run in cells[1].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)


def add_bullets(document: Document, items: list[str], style: str = "List Bullet") -> None:
    for item in items:
        paragraph = document.add_paragraph(item, style=style)
        paragraph.paragraph_format.left_indent = Cm(0.65)
        paragraph.paragraph_format.first_line_indent = Cm(-0.3)


def add_procedure(document: Document, procedure: dict) -> None:
    heading = document.add_heading(procedure["title"], level=3)
    keep_with_next(heading)
    steps_table = document.add_table(rows=0, cols=2)
    steps_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    steps_table.autofit = False
    for index, step in enumerate(procedure["steps"]):
        cells = steps_table.add_row().cells
        cells[0].width = Cm(1)
        cells[1].width = Cm(16)
        cells[0].text = str(index + 1)
        cells[1].text = step
        shade_cell(cells[0], "2C5F7C")
        shade_cell(cells[1], "F6F8FA" if index % 2 == 0 else "FFFFFF")
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=75, start=110, bottom=75, end=110)
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

    outcome_table = document.add_table(rows=1, cols=2)
    outcome_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = (
        ("РЕЗУЛЬТАТ", procedure["result"], "E5F2EA", "337357"),
        ("ЧТО ДАЛЬШЕ", procedure["next"], "E8F2F8", "2C5F7C"),
    )
    for cell, (label, value, fill, color) in zip(outcome_table.rows[0].cells, labels):
        shade_cell(cell, fill)
        set_cell_margins(cell, top=110, start=130, bottom=110, end=130)
        paragraph = cell.paragraphs[0]
        label_run = paragraph.add_run(f"{label}\n")
        label_run.bold = True
        label_run.font.size = Pt(8)
        label_run.font.color.rgb = RGBColor.from_string(color)
        paragraph.add_run(value)


def add_notes(document: Document, notes: list[str]) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, "FFF4CE")
    set_cell_margins(cell, top=120, start=150, bottom=120, end=150)
    title = cell.paragraphs[0]
    title_run = title.add_run("ВАЖНО")
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string("8A5A00")
    for note in notes:
        paragraph = cell.add_paragraph(note, style="List Bullet")
        paragraph.paragraph_format.left_indent = Cm(0.55)


def add_section(document: Document, section: dict, level: int = 2) -> None:
    heading = document.add_heading(section["title"], level=level)
    keep_with_next(heading)
    for paragraph_text in section.get("paragraphs", []):
        paragraph = document.add_paragraph(paragraph_text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if section.get("bullets"):
        add_bullets(document, section["bullets"])
    for procedure in section.get("procedures", []):
        add_procedure(document, procedure)
    if section.get("notes"):
        add_notes(document, section["notes"])


def create_guide(content: dict, target: Path) -> None:
    document = Document()
    configure_document(document)
    add_cover(document, content)

    document.add_heading("Общая часть", level=1)
    document.add_paragraph(
        "Эти правила одинаковы для всех ролей. Названия страниц, кнопок и статусов приведены так, как они показаны в интерфейсе."
    )
    add_journey(document, content["journey"])
    for section in content["common"]:
        add_section(document, section)

    for role_key in ROLE_ORDER:
        role = content["roles"][role_key]
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_heading(f"Работа в роли: {role['label']}", level=1)
        document.add_paragraph(role["intro"])
        add_quick_start(document, role["label"], role["quickStart"], ROLE_COLORS[role_key])
        for section in role["sections"]:
            add_section(document, section)

    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)


def main() -> None:
    content = load_content()
    create_guide(content, OUTPUT_PATH)
    print(f"Created {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
