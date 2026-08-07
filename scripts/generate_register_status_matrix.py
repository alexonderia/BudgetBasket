from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = Path("docs") / "Матрица_статусов_реестра_согласования.docx"

MATRIX_ROWS = [
    {
        "no": "1",
        "item_status": "approved",
        "request_status": "любой",
        "frozen": "—",
        "fixed": "—",
        "cfo_review": "—",
        "cfo_actionable": "—",
        "in_approval": "—",
        "approval_actionable": "—",
        "badge": "Утверждено",
        "caption": "Решение принято",
        "actions": "Нет",
    },
    {
        "no": "2",
        "item_status": "approved_with_changes",
        "request_status": "любой",
        "frozen": "—",
        "fixed": "—",
        "cfo_review": "—",
        "cfo_actionable": "—",
        "in_approval": "—",
        "approval_actionable": "—",
        "badge": "Утверждено с изменениями",
        "caption": "Сумма скорректирована",
        "actions": "Нет",
    },
    {
        "no": "3",
        "item_status": "rejected",
        "request_status": "любой",
        "frozen": "true",
        "fixed": "—",
        "cfo_review": "—",
        "cfo_actionable": "—",
        "in_approval": "—",
        "approval_actionable": "—",
        "badge": "Отклонено",
        "caption": "На доработке",
        "actions": "Нет",
    },
    {
        "no": "4",
        "item_status": "rejected",
        "request_status": "любой",
        "frozen": "false",
        "fixed": "—",
        "cfo_review": "—",
        "cfo_actionable": "—",
        "in_approval": "—",
        "approval_actionable": "—",
        "badge": "Отклонено",
        "caption": "На доработке",
        "actions": "Нет",
    },
    {
        "no": "5",
        "item_status": "on_review",
        "request_status": "draft",
        "frozen": "—",
        "fixed": "—",
        "cfo_review": "—",
        "cfo_actionable": "—",
        "in_approval": "—",
        "approval_actionable": "—",
        "badge": "Черновик",
        "caption": "Не отправлено",
        "actions": "Нет",
    },
    {
        "no": "6",
        "item_status": "on_review",
        "request_status": "cancelled",
        "frozen": "—",
        "fixed": "—",
        "cfo_review": "—",
        "cfo_actionable": "—",
        "in_approval": "—",
        "approval_actionable": "—",
        "badge": "Заявка отменена",
        "caption": "Отменена",
        "actions": "Нет",
    },
    {
        "no": "7",
        "item_status": "on_review",
        "request_status": "rejected",
        "frozen": "—",
        "fixed": "—",
        "cfo_review": "—",
        "cfo_actionable": "—",
        "in_approval": "—",
        "approval_actionable": "—",
        "badge": "Заявка отклонена",
        "caption": "Заявка отклонена",
        "actions": "Нет",
    },
    {
        "no": "8",
        "item_status": "on_review",
        "request_status": "on_review",
        "frozen": "—",
        "fixed": "—",
        "cfo_review": "да",
        "cfo_actionable": "да",
        "in_approval": "—",
        "approval_actionable": "—",
        "badge": "Ожидает вашего решения",
        "caption": "Можно принять решение",
        "actions": "Да (проверка ЦФО)",
    },
    {
        "no": "9",
        "item_status": "on_review",
        "request_status": "on_review",
        "frozen": "—",
        "fixed": "—",
        "cfo_review": "да",
        "cfo_actionable": "нет",
        "in_approval": "—",
        "approval_actionable": "—",
        "badge": "Проверка ЦФО",
        "caption": "Ждёт ответственного ЦФО",
        "actions": "Нет",
    },
    {
        "no": "10",
        "item_status": "on_review",
        "request_status": "approved*",
        "frozen": "—",
        "fixed": "да",
        "cfo_review": "—",
        "cfo_actionable": "—",
        "in_approval": "да",
        "approval_actionable": "—",
        "badge": "Зафиксировано",
        "caption": "Зафиксировано",
        "actions": "Нет",
    },
    {
        "no": "11",
        "item_status": "on_review",
        "request_status": "approved*",
        "frozen": "да",
        "fixed": "нет",
        "cfo_review": "—",
        "cfo_actionable": "—",
        "in_approval": "да",
        "approval_actionable": "—",
        "badge": "На доработке",
        "caption": "Ожидает исправлений",
        "actions": "Нет",
    },
    {
        "no": "12",
        "item_status": "on_review",
        "request_status": "approved*",
        "frozen": "нет",
        "fixed": "нет",
        "cfo_review": "—",
        "cfo_actionable": "—",
        "in_approval": "да",
        "approval_actionable": "да",
        "badge": "Ожидает вашего решения",
        "caption": "Можно принять решение",
        "actions": "Да (маршрут согласования)",
    },
    {
        "no": "13",
        "item_status": "on_review",
        "request_status": "approved*",
        "frozen": "нет",
        "fixed": "нет",
        "cfo_review": "—",
        "cfo_actionable": "—",
        "in_approval": "да",
        "approval_actionable": "нет",
        "badge": "Ожидает предыдущих этапов",
        "caption": "Ждёт предыдущих этапов",
        "actions": "Нет",
    },
    {
        "no": "14",
        "item_status": "on_review",
        "request_status": "прочее",
        "frozen": "—",
        "fixed": "—",
        "cfo_review": "нет",
        "cfo_actionable": "—",
        "in_approval": "нет",
        "approval_actionable": "—",
        "badge": "На рассмотрении",
        "caption": "Ожидает решения",
        "actions": "Нет",
    },
    {
        "no": "15",
        "item_status": "прочее",
        "request_status": "—",
        "frozen": "—",
        "fixed": "—",
        "cfo_review": "—",
        "cfo_actionable": "—",
        "in_approval": "—",
        "approval_actionable": "—",
        "badge": "Не начато",
        "caption": "Не отправлено",
        "actions": "Нет",
    },
]

HEADERS = [
    "№",
    "req_items.status",
    "requests.status",
    "frozen",
    "fixed",
    "is_cfo_review",
    "is_cfo_review_actionable",
    "is_in_approval",
    "is_approval_actionable",
    "Бейдж",
    "Подпись",
    "Действия",
]


def shade(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    properties.append(shading)


def set_cell_text(cell, value: str, *, bold: bool = False, size: int = 8) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(value)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def add_paragraph(doc: Document, text: str, *, bold: bool = False) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    for style_name in ("Normal", "Title", "Heading 1", "Heading 2"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        if style_name == "Title":
            style.font.color.rgb = RGBColor(31, 78, 121)

    title = doc.add_heading("BudgetBasket", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Матрица: комбинация состояний → колонка «Статус»")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.bold = True
        run.font.size = Pt(12)

    add_paragraph(
        doc,
        "Документ описывает отображение статуса строк бюджета (req_item) в реестре согласования. "
        "Проверка условий выполняется сверху вниз: применяется первая подходящая строка матрицы.",
    )
    add_paragraph(
        doc,
        "Источник логики: frontend/src/components/approval-register/registryConfig.ts → rowRegistryStatus().",
        bold=True,
    )
    add_paragraph(
        doc,
        "Флаги is_cfo_review_actionable и is_approval_actionable зависят от роли и назначений текущего пользователя.",
    )

    doc.add_heading("Матрица строк бюджета", level=1)

    table = doc.add_table(rows=1, cols=len(HEADERS))
    table.style = "Table Grid"
    table.autofit = False

    for index, header in enumerate(HEADERS):
        cell = table.rows[0].cells[index]
        set_cell_text(cell, header, bold=True, size=8)
        shade(cell, "D9EAF7")

    for row in MATRIX_ROWS:
        values = [
            row["no"],
            row["item_status"],
            row["request_status"],
            row["frozen"],
            row["fixed"],
            row["cfo_review"],
            row["cfo_actionable"],
            row["in_approval"],
            row["approval_actionable"],
            row["badge"],
            row["caption"],
            row["actions"],
        ]
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_cell_text(cells[index], value, bold=index >= 9, size=8)
            if index >= 9:
                shade(cells[index], "F8FBFF")

    doc.add_heading("Примечания", level=1)
    notes = [
        "Строки с req_items.status = deleted в реестр не попадают.",
        "После проверки ЦФО заявка переходит в requests.status = approved, а строки продолжают движение по маршруту позиции (is_in_approval).",
        "Поле approval_stage заполняется по текущему шагу маршрута: «Проверка экономистом ЦФО», «Согласование проверяющим», «Финальное согласование ЗГД».",
        "Действия доступны, если is_cfo_review_actionable = да или (is_approval_actionable = да и есть position_id).",
        "Редактирование суммы «Согласовано» доступно ролям employee и economist только для actionable-строк.",
        "При наведении на статус в таблице показывается полная подсказка (hint); под бейджем — краткая подпись (shortHint).",
    ]
    for note in notes:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(note)
        run.font.size = Pt(10)
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build_document().save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
